#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Convert Markdown to AGIC-styled DOCX using python-docx and mistune.

This is the new Option A converter that uses:
- mistune (v3.x) for robust markdown parsing
- python-docx for Word document generation with AGIC template

Usage:
    python markdown_to_docx_v2.py input.md output.docx [--template template.docx]
"""

import argparse
import base64
import os
import re
import shutil
import sys
import tempfile
import urllib.request
import zipfile
from datetime import date as _date
from pathlib import Path
from typing import List, Dict, Any, Optional

import mistune
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Force UTF-8 output
if sys.stdout.encoding != 'utf-8':
    sys.stdout.reconfigure(encoding='utf-8')


# AGIC Style Mapping
AGIC_STYLES = {
    'heading1': 'AGIC Titolo Copertina',  # H1 → cover page title style
    'heading2': 'AGIC Titolo Paragrafo',
    'heading3': 'AGIC Titolo 2',
    'heading4': 'AGIC Titolo 3',
    'heading5': None,               # H5 → bold paragraph (AGIC corpo + bold)
    'heading6': None,               # H6 → bold paragraph (AGIC corpo + bold)
    'paragraph': 'AGIC corpo',
    'code': 'AGIC corpo',           # Code blocks use corpo with monospace
    'list_ordered': 'AGIC elenco numerato',
    'list_unordered': 'AGIC elenco puntato',
    'table': 'AGICTable1',
}



def convert_markdown_to_docx(
    markdown_path: Path,
    output_path: Path,
    template_path: Path,
    organization: str = '',
    date: str = '',
) -> None:
    """Convert Markdown file to AGIC-styled DOCX.
    
    Args:
        markdown_path: Path to input markdown file
        output_path: Path to output DOCX file
        template_path: Path to AGIC template DOCX
        organization: Client/organization name for the Cliente field.
                      Overrides YAML frontmatter ``organization:`` if provided.
        date: Publication date (dd/MM/yyyy) for Data pubblicazione.
              Overrides YAML frontmatter ``date:`` if provided.
              Defaults to today when neither this nor YAML supplies a value.
    """
    print(f"Converting {markdown_path} to {output_path}")
    
    # Read markdown
    with open(markdown_path, 'r', encoding='utf-8') as f:
        markdown_text = f.read()
    
    # Extract YAML frontmatter metadata
    metadata = {}
    if markdown_text.startswith('---'):
        parts = markdown_text.split('---', 2)
        if len(parts) >= 3:
            yaml_text = parts[1].strip()
            markdown_text = parts[2].strip()
            
            # Parse YAML manually (simple key: value format)
            for line in yaml_text.split('\n'):
                if ':' in line:
                    key, value = line.split(':', 1)
                    metadata[key.strip()] = value.strip()
            
            print(f"  Extracted metadata: {list(metadata.keys())}")
    
    # Load AGIC template
    doc = Document(str(template_path))
    print(f"  Loaded AGIC template: {template_path}")
    
    # Populate cover page with metadata if available
    if metadata:
        _populate_cover_page(doc, metadata)
        print("  Populated cover page fields")
    
    # Parse markdown using mistune to get AST
    md = mistune.create_markdown(renderer='ast', plugins=['table', 'strikethrough'])
    tokens = md(markdown_text)
    
    # Render tokens to document
    _render_tokens_to_doc(doc, tokens, markdown_path)
    
    # Save document
    doc.save(str(output_path))

    # Post-process: update data-bound custom XML parts that Word refreshes on open.
    # These cannot be updated via python-docx alone — requires zip-level edits.
    # CLI args take precedence over YAML frontmatter; YAML overrides empty defaults.
    publish_date = date or metadata.get('date', '') or _date.today().strftime('%d/%m/%Y')
    company = organization or metadata.get('organization', '')
    _update_docx_custom_xml_parts(output_path, publish_date=publish_date, company=company)

    print(f"✓ Conversion complete: {output_path}")
    print(f"  Size: {output_path.stat().st_size / 1024:.1f} KB")


def _update_docx_custom_xml_parts(docx_path: Path, publish_date: str = '', company: str = '') -> None:
    """Zip-level post-processing: update data-bound custom XML parts.

    The AGIC template SDTs with w:dataBinding ignore w:sdtContent on file open
    and refresh from their bound XML source instead. We must update those sources:
    - customXml/item1.xml  → CoverPageProperties.PublishDate  (Data pubblicazione)
    - docProps/app.xml     → Properties.Company               (Cliente)
    """
    from xml.sax.saxutils import escape as _xml_escape

    if not publish_date and not company:
        return

    tmp_fd, tmp_path_str = tempfile.mkstemp(suffix='.docx')
    os.close(tmp_fd)
    tmp = Path(tmp_path_str)
    try:
        with zipfile.ZipFile(str(docx_path), 'r') as zin:
            with zipfile.ZipFile(str(tmp), 'w', compression=zipfile.ZIP_DEFLATED) as zout:
                for name in zin.namelist():
                    data = zin.read(name)

                    if publish_date and name == 'customXml/item1.xml':
                        text = data.decode('utf-8')
                        text = re.sub(
                            r'<PublishDate>[^<]*</PublishDate>',
                            f'<PublishDate>{_xml_escape(publish_date)}</PublishDate>',
                            text,
                        )
                        data = text.encode('utf-8')

                    elif company and name == 'docProps/app.xml':
                        text = data.decode('utf-8')
                        escaped = _xml_escape(company)
                        if re.search(r'<Company>', text):
                            text = re.sub(
                                r'<Company>[^<]*</Company>',
                                f'<Company>{escaped}</Company>',
                                text,
                            )
                        else:
                            text = text.replace(
                                '</Properties>',
                                f'<Company>{escaped}</Company></Properties>',
                            )
                        data = text.encode('utf-8')

                    zout.writestr(name, data)

        shutil.move(str(tmp), str(docx_path))
    except Exception:
        if tmp.exists():
            tmp.unlink()
        raise


def _populate_cover_page(doc: Document, metadata: Dict[str, str]) -> None:
    """Populate AGIC cover page content controls with metadata.
    
    Args:
        doc: The document object
        metadata: Dictionary with organization, project, date, title, subtitle
    """
    # Extract and map metadata fields
    title = metadata.get('title', '')
    subtitle = metadata.get('subtitle', '')
    
    # If no explicit title, use project field
    if not title:
        project = metadata.get('project', '')
        if ' - ' in project:
            # Split "Type - Name" into subtitle and title
            parts = project.split(' - ', 1)
            subtitle = subtitle or parts[0]
            title = parts[1]
        else:
            title = project
    
    organization = metadata.get('organization', '')
    date = metadata.get('date', '')
    
    # Update document core properties (for data binding)
    core_props = doc.core_properties
    if title:
        core_props.title = title
    if organization:
        core_props.author = organization  # Or use company if available
    # Note: core_properties doesn't have a direct 'company' field
    # We'll need to update the app.xml separately
    
    # Mapping: control alias → value
    control_mappings = {
        'Titolo': title,
        'Cliente': organization,
        'Data pubblicazione': date
    }
    
    # Update content controls in document body
    _update_content_controls(doc._element, control_mappings)
    
    # Update content controls in all headers
    for section in doc.sections:
        if section.header:
            _update_content_controls(section.header._element, control_mappings)
    
    # Update subtitle (non-control, static text)
    if subtitle:
        _update_subtitle(doc, subtitle)


def _update_content_controls(element, mappings: Dict[str, str]) -> None:
    """Update all content controls in an element by alias.
    
    Args:
        element: XML element to search (document body or header)
        mappings: Dictionary mapping control alias to new value
    """
    # Find all structured document tags (content controls)
    # Use namespace-prefixed XPath directly
    for sdt in element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sdt'):
        # Get the alias of this control
        alias_nodes = sdt.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}alias')
        if not alias_nodes:
            continue
            
        alias_value = alias_nodes[0].get(qn('w:val'))
        
        # Check if we have a value for this control
        if alias_value in mappings and mappings[alias_value]:
            _set_content_control_text(sdt, mappings[alias_value])


def _set_content_control_text(sdt, text: str) -> None:
    """Set the text content of a content control.
    
    Args:
        sdt: The w:sdt element
        text: New text value
    """
    w_ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    
    # Find sdtContent
    content_nodes = sdt.findall(f'.//{{{w_ns}}}sdtContent')
    if not content_nodes:
        return
    
    sdt_content = content_nodes[0]
    
    # Find all text nodes within the content
    text_nodes = sdt_content.findall(f'.//{{{w_ns}}}t')
    
    if text_nodes:
        # Replace text in first text node, clear others
        text_nodes[0].text = text
        for node in text_nodes[1:]:
            node.text = ''
    else:
        # No text node exists, need to create the structure
        # Find or create run
        runs = sdt_content.findall(f'.//{{{w_ns}}}r')
        if runs:
            run = runs[0]
        else:
            # Create paragraph if needed
            paras = sdt_content.findall(f'.//{{{w_ns}}}p')
            if not paras:
                para = OxmlElement('w:p')
                sdt_content.append(para)
                paras = [para]
            
            # Create run
            run = OxmlElement('w:r')
            paras[0].append(run)
        
        # Create text element
        t = OxmlElement('w:t')
        t.text = text
        run.append(t)


def _update_subtitle(doc: Document, subtitle: str) -> None:
    """Update subtitle text (non-content-control).
    
    The subtitle is static text with style 'SottotitolocopertinaCarattere'.
    
    Args:
        doc: The document object
        subtitle: New subtitle text
    """
    w_ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    
    # Search for runs with the subtitle character style
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            # Check if run has the subtitle style
            style_nodes = run._element.findall(f'.//{{{w_ns}}}rStyle')
            for style_node in style_nodes:
                if style_node.get(qn('w:val')) == 'SottotitolocopertinaCarattere':
                    run.text = subtitle
                    return  # Found and updated, done


def _apply_cover_title_style(doc: Document, style_name: str) -> None:
    """Apply a paragraph style to the cover-page paragraph that hosts the Titolo SDT."""
    w_ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    for paragraph in doc.paragraphs:
        for sdt in paragraph._element.findall(f'{{{w_ns}}}sdt'):
            alias_nodes = sdt.findall(f'.//{{{w_ns}}}alias')
            if alias_nodes and alias_nodes[0].get(qn('w:val')) == 'Titolo':
                paragraph.style = style_name
                return


def _render_tokens_to_doc(doc: Document, tokens: List[Dict[str, Any]], source_file: Path = None) -> None:
    """Render mistune tokens to Word document."""
    # First pass: extract H1 as document title → populate cover page title
    # control and apply AGICTitoloCopertina paragraph style.
    for token in tokens:
        if token.get('type') == 'heading' and token.get('attrs', {}).get('level', 1) == 1:
            h1_text = _extract_text_from_tokens(token.get('children', []))
            if h1_text:
                # The Titolo SDT is data-bound to core_properties.title;
                # set both so Word's data binding AND direct XML readers show the title.
                doc.core_properties.title = h1_text
                _update_content_controls(doc._element, {'Titolo': h1_text})
                _apply_cover_title_style(doc, AGIC_STYLES['heading1'])
                for section in doc.sections:
                    if section.header:
                        _update_content_controls(section.header._element, {'Titolo': h1_text})
            break

    for token in tokens:
        token_type = token.get('type')
        
        if token_type == 'heading':
            _render_heading(doc, token)
        elif token_type == 'paragraph':
            _render_paragraph(doc, token, source_file)
        elif token_type == 'block_code':
            _render_code_block(doc, token, source_file)
        elif token_type == 'list':
            _render_list(doc, token)
        elif token_type == 'table':
            _render_table(doc, token)
        elif token_type == 'block_image':
            _render_image(doc, token, source_file)
        elif token_type == 'thematic_break':
            doc.add_paragraph()  # Empty paragraph for HR
        elif token_type == 'blank_line':
            pass  # Skip blank lines


def _clean_heading_runs(paragraph) -> None:
    """Strip anchor tags and leading manual numbering from heading runs."""
    first_run = True
    for run in paragraph.runs:
        if run.text:
            run.text = re.sub(r'\s*\{#[-_A-Za-z0-9]+\}\s*$', '', run.text)
            if first_run:
                run.text = re.sub(r'^\s*\d+(\.\d+)*\.?\s+', '', run.text)
                first_run = False


def _render_heading(doc: Document, token: Dict[str, Any]) -> None:
    """Render heading token.

    H1 → cover page title (content control already populated in
    _render_tokens_to_doc); not rendered as body paragraph.
    H2–H4 → AGIC heading styles.
    H5–H6 → bold body paragraph (AGIC corpo with bold runs).
    """
    level = token.get('attrs', {}).get('level', 1)

    # H1 text is injected into the cover-page Titolo content control
    # in _render_tokens_to_doc; skip body rendering.
    if level == 1:
        return

    children = token.get('children', [])
    style_key = f'heading{level}'
    style = AGIC_STYLES.get(style_key)

    # H5/H6: render as bold AGIC corpo paragraph (no heading style)
    if style is None:
        p = doc.add_paragraph()
        p.style = AGIC_STYLES['paragraph']
        _add_formatted_text(p, '', children)
        _clean_heading_runs(p)
        for run in p.runs:
            run.bold = True
        return

    p = doc.add_paragraph()
    p.style = style

    # Add formatted text from children
    _add_formatted_text(p, '', children)

    # Clean anchor tags and strip leading numbering (e.g. "5.", "5.1", "5.1.2.")
    # AGIC template handles auto-numbering; leading numbers in markdown must be removed.
    _clean_heading_runs(p)


def _render_paragraph(doc: Document, token: Dict[str, Any], source_file: Path = None) -> None:
    """Render paragraph token."""
    children = token.get('children', [])
    
    if not children:
        return
    
    # Check if paragraph contains only an image - render as block image
    if len(children) == 1 and children[0].get('type') == 'image':
        _render_image(doc, children[0], source_file)
        return
    
    p = doc.add_paragraph()
    p.style = AGIC_STYLES['paragraph']
    
    text = _extract_text_from_tokens(children)
    _add_formatted_text(p, text, children, source_file)


def _render_mermaid(doc: Document, code: str, source_file: Path = None) -> None:
    """Render a mermaid diagram block into the document.

    Rendering tiers (tried in order):
      1. mermaid.ink public API — stdlib urllib + base64, zero extra deps, needs internet
      2. mmdc (mermaid-cli) via mermaid_renderer.py — offline, needs Node + mmdc installed
      3. Styled code block fallback — always works, never loses diagram source
    """
    png_bytes: Optional[bytes] = None

    # Tier 1: mermaid.ink HTTP API
    try:
        encoded = base64.urlsafe_b64encode(code.encode('utf-8')).decode('ascii')
        url = f"https://mermaid.ink/img/{encoded}"
        req = urllib.request.Request(url, headers={"User-Agent": "agic-docx/1.0"})
        with urllib.request.urlopen(req, timeout=10) as resp:
            if resp.status == 200:
                png_bytes = resp.read()
    except Exception:
        pass  # fall through to Tier 2

    # Tier 2: local mmdc via mermaid_renderer.py
    if not png_bytes:
        try:
            import sys as _sys
            _scripts_dir = str(Path(__file__).parent)
            if _scripts_dir not in _sys.path:
                _sys.path.insert(0, _scripts_dir)
            from mermaid_renderer import MermaidRenderer
            renderer = MermaidRenderer()
            if renderer.is_available():
                png_bytes = renderer.render(code)
        except Exception:
            pass  # fall through to Tier 3

    # Tiers 1+2 succeeded — embed as image
    if png_bytes:
        import io
        img_stream = io.BytesIO(png_bytes)
        p = doc.add_paragraph()
        p.style = AGIC_STYLES['paragraph']
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        run = p.add_run()
        run.add_picture(img_stream, width=Inches(5.5))
        return

    # Tier 3: fallback — show source as styled code block with label
    label = doc.add_paragraph("[Mermaid Diagram — rendered at open time or paste source below]")
    label.style = AGIC_STYLES['paragraph']
    label.runs[0].italic = True

    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls
    for line in code.rstrip('\n').split('\n'):
        p = doc.add_paragraph()
        p.style = AGIC_STYLES['code']
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        shading_elm = parse_xml(r'<w:shd {} w:fill="FFF8DC"/>'.format(nsdecls('w')))  # pale yellow
        p._element.get_or_add_pPr().append(shading_elm)
        run = p.add_run(line if line else ' ')
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        for r in p._element.xpath('.//w:r'):
            for t in r.xpath('.//w:t'):
                t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')



def _render_code_block(doc: Document, token: Dict[str, Any], source_file: Path = None) -> None:
    """Render code block token with proper formatting."""
    code = token.get('raw', '')
    info = token.get('attrs', {}).get('info', '')

    # Mermaid diagrams — use 3-tier renderer
    if info and 'mermaid' in info.lower():
        _render_mermaid(doc, code, source_file)
        return

    # Remove trailing newline if present
    if code.endswith('\n'):
        code = code[:-1]
    
    lines = code.split('\n')
    
    # Create one paragraph per line to preserve line breaks and spacing
    for line in lines:
        p = doc.add_paragraph()
        p.style = AGIC_STYLES['code']
        
        # Set left alignment (not justified)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        
        # Add gray background shading
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls
        shading_elm = parse_xml(r'<w:shd {} w:fill="F0F0F0"/>'.format(nsdecls('w')))
        p._element.get_or_add_pPr().append(shading_elm)
        
        # Add the line with preserved spaces
        run = p.add_run(line if line else ' ')  # Empty lines get a space
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        
        # Preserve spaces using XML space attribute
        if line:
            for r in p._element.xpath('.//w:r'):
                for t in r.xpath('.//w:t'):
                    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')


def _render_list(doc: Document, token: Dict[str, Any]) -> None:
    """Render list token using AGIC list styles."""
    children = token.get('children', [])
    ordered = token.get('attrs', {}).get('ordered', False)
    style = AGIC_STYLES['list_ordered'] if ordered else AGIC_STYLES['list_unordered']
    
    for item in children:
        if item.get('type') == 'list_item':
            item_children = item.get('children', [])
            text = _extract_text_from_tokens(item_children)
            
            p = doc.add_paragraph()
            p.style = style
            p.add_run(text)


def _add_hyperlink(paragraph, text: str, url: str) -> None:
    """Add a hyperlink to a paragraph.
    
    Args:
        paragraph: The paragraph to add the hyperlink to
        text: The text to display
        url: The URL to link to
    """
    # Create relationship for external link
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    
    # Create the hyperlink element
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    
    # Create a run element within the hyperlink
    run = OxmlElement('w:r')
    
    # Add run properties for hyperlink styling
    rPr = OxmlElement('w:rPr')
    
    # Add color (blue)
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0563C1')
    rPr.append(color)
    
    # Add underline
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    
    run.append(rPr)
    
    # Add the text
    t = OxmlElement('w:t')
    t.text = text
    run.append(t)
    
    hyperlink.append(run)
    paragraph._element.append(hyperlink)


def _render_image(doc: Document, token: Dict[str, Any], source_file: Path = None) -> None:
    """Render image token - try to embed, fallback to placeholder.
    
    Args:
        doc: The document object
        token: The image token
        source_file: Path to source markdown file (for relative paths)
    """
    alt_text = token.get('attrs', {}).get('alt', 'Image')
    url = token.get('attrs', {}).get('url', '')
    
    p = doc.add_paragraph()
    p.style = AGIC_STYLES['paragraph']
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    image_inserted = False
    
    # Try to insert local images
    if url and not url.startswith('http://') and not url.startswith('https://'):
        image_path = Path(url)
        
        # Make relative to markdown file location if not absolute
        if not image_path.is_absolute() and source_file:
            image_path = source_file.parent / image_path
        
        # Try to insert the image
        if image_path.exists() and image_path.is_file():
            try:
                # Insert image with max width of 6 inches
                run = p.add_run()
                run.add_picture(str(image_path), width=Inches(6))
                image_inserted = True
                
                # Add caption below
                if alt_text and alt_text != 'Image':
                    caption_p = doc.add_paragraph()
                    caption_p.style = AGIC_STYLES['paragraph']
                    caption_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    caption_run = caption_p.add_run(alt_text)
                    caption_run.font.size = Pt(9)
                    caption_run.italic = True
                    caption_run.font.color.rgb = RGBColor(96, 96, 96)
            except Exception as e:
                # Failed to insert, will show placeholder
                pass
    
    # Fallback to placeholder if image not inserted
    if not image_inserted:
        if url.startswith('http://') or url.startswith('https://'):
            # External image - show placeholder with URL
            run = p.add_run(f"[External Image: {alt_text}]")
            run.italic = True
            run.font.color.rgb = RGBColor(128, 128, 128)
            
            # Add URL on next line
            p2 = doc.add_paragraph()
            p2.style = AGIC_STYLES['paragraph']
            p2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run2 = p2.add_run(f"{url}")
            run2.font.size = Pt(8)
            run2.font.color.rgb = RGBColor(128, 128, 128)
        else:
            # Local image not found - show placeholder
            run = p.add_run(f"[Image not found: {alt_text}]")
            run.italic = True
            run.font.color.rgb = RGBColor(192, 0, 0)  # Red for missing
            
            if url:
                p2 = doc.add_paragraph()
                p2.style = AGIC_STYLES['paragraph']
                p2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run2 = p2.add_run(f"Path: {url}")
                run2.font.size = Pt(8)
                run2.font.color.rgb = RGBColor(192, 0, 0)


def _extract_text_from_tokens(tokens: List[Dict[str, Any]]) -> str:
    """Extract plain text from tokens."""
    result = []
    for token in tokens:
        if token.get('type') == 'text':
            result.append(token.get('raw', ''))
        elif token.get('type') in ('strong', 'emphasis', 'codespan', 'link'):
            # Recurse into children
            children = token.get('children', [])
            if children:
                result.append(_extract_text_from_tokens(children))
            else:
                result.append(token.get('raw', ''))
        elif 'children' in token:
            result.append(_extract_text_from_tokens(token['children']))
    return ''.join(result)


def _add_formatted_text(paragraph, text: str, tokens: List[Dict[str, Any]], source_file: Path = None) -> None:
    """Add text to paragraph with inline formatting."""
    for token in tokens:
        token_type = token.get('type')
        
        if token_type == 'text':
            raw = token.get('raw', '')
            paragraph.add_run(raw)
        elif token_type == 'strong':
            # Bold text - extract from children
            children = token.get('children', [])
            text_content = _extract_text_from_tokens(children)
            run = paragraph.add_run(text_content)
            run.bold = True
        elif token_type == 'emphasis':
            # Italic text - extract from children
            children = token.get('children', [])
            text_content = _extract_text_from_tokens(children)
            run = paragraph.add_run(text_content)
            run.italic = True
        elif token_type == 'codespan':
            # Inline code — monospace font + light gray background
            raw = token.get('raw', '')
            run = paragraph.add_run(raw)
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            # Add light gray background shading (w:shd on rPr)
            rPr = run._element.get_or_add_rPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), 'F0F0F0')
            rPr.append(shd)
        elif token_type == 'link':
            # Add hyperlink
            children = token.get('children', [])
            link_text = _extract_text_from_tokens(children)
            link_url = token.get('attrs', {}).get('url', '')
            _add_hyperlink(paragraph, link_text, link_url)
        elif token_type == 'image':
            # Inline image - add placeholder text
            alt_text = token.get('attrs', {}).get('alt', 'Image')
            url = token.get('attrs', {}).get('url', '')
            run = paragraph.add_run(f"[Image: {alt_text}]")
            run.italic = True
        elif 'children' in token:
            # Recurse into other tokens with children
            _add_formatted_text(paragraph, '', token['children'], source_file)


def _render_table(doc: Document, token: Dict[str, Any]) -> None:
    """Render table token."""
    children = token.get('children', [])
    
    # Extract table structure
    header_cells = []
    body_rows = []
    
    for child in children:
        if child.get('type') == 'table_head':
            # Extract header cells
            for cell in child.get('children', []):
                if cell.get('type') == 'table_cell':
                    cell_children = cell.get('children', [])
                    cell_text = _extract_text_from_tokens(cell_children)
                    header_cells.append(cell_text)
        
        elif child.get('type') == 'table_body':
            # Extract body rows
            for row in child.get('children', []):
                if row.get('type') == 'table_row':
                    row_cells = []
                    for cell in row.get('children', []):
                        if cell.get('type') == 'table_cell':
                            cell_children = cell.get('children', [])
                            cell_text = _extract_text_from_tokens(cell_children)
                            # Check if cell has formatting or links
                            has_formatting = any(c.get('type') in ('strong', 'emphasis', 'codespan', 'link', 'image') 
                                                for c in cell_children)
                            row_cells.append((cell_text, cell_children, has_formatting))
                    body_rows.append(row_cells)
    
    if not header_cells or not body_rows:
        return
    
    # Create Word table
    table = doc.add_table(rows=len(body_rows) + 1, cols=len(header_cells))
    table.style = AGIC_STYLES['table']
    
    # Header row
    for col_idx, header_text in enumerate(header_cells):
        cell = table.rows[0].cells[col_idx]
        cell.text = header_text
        # Bold header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    # Body rows
    for row_idx, row_data in enumerate(body_rows):
        for col_idx, (cell_text, cell_tokens, has_formatting) in enumerate(row_data):
            if col_idx < len(table.rows[row_idx + 1].cells):
                cell = table.rows[row_idx + 1].cells[col_idx]
                if has_formatting and cell_tokens:
                    # Clear default paragraph and add formatted text
                    cell.text = ''
                    p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
                    _add_formatted_text(p, cell_text, cell_tokens)
                else:
                    cell.text = cell_text




def main():
    """Main entry point."""
    parser = argparse.ArgumentParser(
        description="Convert Markdown to AGIC-styled DOCX using python-docx + mistune"
    )
    parser.add_argument("input", type=Path, help="Input markdown file")
    parser.add_argument("output", type=Path, help="Output DOCX file")
    parser.add_argument(
        "--template",
        type=Path,
        default=Path(__file__).parent.parent / "assets" / "agic-template.docx",
        help="AGIC template DOCX (default: ../assets/agic-template.docx)"
    )
    parser.add_argument(
        "--organization",
        default='',
        help="Client/organization name for the Cliente field (overrides YAML frontmatter)"
    )
    parser.add_argument(
        "--date",
        default='',
        help="Publication date dd/MM/yyyy for Data pubblicazione (overrides YAML; defaults to today)"
    )
    
    args = parser.parse_args()
    
    if not args.input.exists():
        print(f"Error: Input file not found: {args.input}", file=sys.stderr)
        sys.exit(1)
    
    if not args.template.exists():
        print(f"Error: Template not found: {args.template}", file=sys.stderr)
        sys.exit(1)
    
    try:
        convert_markdown_to_docx(
            args.input, args.output, args.template,
            organization=args.organization,
            date=args.date,
        )
    except Exception as e:
        print(f"Error during conversion: {e}", file=sys.stderr)
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == "__main__":
    main()
